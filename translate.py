import streamlit as st
import requests
from docx import Document
from io import BytesIO
import json
import time

# Configuration
API_URL = "https://bagwkqqw6a6i3e7p.us-east-1.aws.endpoints.huggingface.cloud"
HF_TOKEN = st.secrets["huggingface"]["api_key"]
headers = {
    "Authorization": f"Bearer {HF_TOKEN}",
    "Content-Type": "application/json"
}

# Language mapping with better coverage
INDIC_LANGS = {
    "English": "eng_Latn",
    "Hindi": "hin_Deva",
    "Marathi": "mar_Deva", 
    "Bengali": "ben_Beng",
    "Tamil": "tam_Taml",
    "Telugu": "tel_Telu",
    "Gujarati": "guj_Gujr",
    "Kannada": "kan_Knda",
    "Malayalam": "mal_Mlym",
    "Punjabi": "pan_Guru",
    "Urdu": "urd_Arab",
    "Odia": "ory_Orya",
    "Assamese": "asm_Beng",
    "Nepali": "nep_Deva"
}

# Style prompts for natural translation
STYLE_PROMPTS = {
    "formal": "",
    "conversational": "Translate in a natural, conversational style as people speak in daily life. Avoid overly formal or literary language.",
    "casual": "Translate in a casual, friendly tone using common everyday words that people use in normal conversation.",
    "simple": "Use simple, easy-to-understand words that are commonly used in everyday speech."
}

def preprocess_text(text, style="conversational"):
    """Add style instruction to improve translation quality"""
    if style != "formal" and text.strip():
        instruction = STYLE_PROMPTS[style]
        return f"{instruction}\n\nText to translate: {text}"
    return text

def translate_text(text, source_lang, target_lang, style="conversational", max_retries=3):
    """Translate text with retry logic and style enhancement"""
    if not text.strip():
        return text
    
    # Preprocess for natural translation
    processed_text = preprocess_text(text, style)
    
    payload = {
        "inputs": processed_text,
        "parameters": {
            "src_lang": source_lang,
            "tgt_lang": target_lang,
            "max_length": 512,
            "temperature": 0.7,  # Add some creativity
            "do_sample": True
        }
    }
    
    for attempt in range(max_retries):
        try:
            response = requests.post(
                API_URL, 
                headers=headers, 
                data=json.dumps(payload), 
                timeout=45
            )
            response.raise_for_status()
            result = response.json()
            
            if isinstance(result, list) and len(result) > 0:
                translation = result[0].get("translation_text", "")
                # Clean up the translation
                translation = clean_translation(translation, style)
                return translation
            else:
                return f"[ERROR] Unexpected response format"
                
        except requests.exceptions.Timeout:
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)  # Exponential backoff
                continue
            return f"[ERROR] Translation timeout after {max_retries} attempts"
        except requests.exceptions.RequestException as e:
            return f"[ERROR] Network error: {str(e)}"
        except Exception as e:
            return f"[ERROR] {str(e)}"
    
    return "[ERROR] Translation failed after all retries"

def clean_translation(translation, style):
    """Clean and post-process translation"""
    if not translation:
        return translation
    
    # Remove style instruction if it appears in translation
    for prompt in STYLE_PROMPTS.values():
        if prompt and prompt in translation:
            translation = translation.replace(prompt, "")
    
    # Remove common prefixes that might appear
    prefixes_to_remove = [
        "Text to translate:",
        "Translation:",
        "Translated text:",
        "Output:"
    ]
    
    for prefix in prefixes_to_remove:
        if translation.startswith(prefix):
            translation = translation[len(prefix):].strip()
    
    return translation.strip()

def extract_text_from_docx(uploaded_file):
    """Extract text from docx file with better structure preservation"""
    doc = Document(uploaded_file)
    content = []
    
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            para_text = ""
            for run in element:
                if run.text:
                    para_text += run.text
            if para_text.strip():
                content.append(("paragraph", para_text.strip()))
        elif element.tag.endswith('tbl'):  # Table
            content.append(("table", "[Table content - manual review needed]"))
    
    return content

# Streamlit UI
st.set_page_config(
    page_title="Document Translator",
    page_icon="🌐",
    layout="wide"
)

st.title("Indian Language Document Translator")
st.write("Translate Word documents between Indian languages using Meta's NLLB-200 model")

# Create columns for better layout
col1, col2 = st.columns([1, 2])

with col1:
    st.subheader("Settings")
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose a Word document", 
        type=["docx"],
        help="Upload a .docx file to translate"
    )
    
    # Language selection
    source = st.selectbox(
        "From Language", 
        list(INDIC_LANGS.keys()), 
        index=0,
        help="Select the source language of your document"
    )
    
    target = st.selectbox(
        "To Language", 
        list(INDIC_LANGS.keys()), 
        index=1,
        help="Select the target language for translation"
    )
    
    # Translation style
    style = st.selectbox(
        "Translation Style",
        ["conversational", "casual", "simple", "formal"],
        index=0,
        help="Choose the style of translation"
    )
    
    # Batch size for processing
    batch_size = st.slider(
        "Batch Size", 
        min_value=1, 
        max_value=10, 
        value=3,
        help="Number of paragraphs to process at once"
    )

with col2:
    st.subheader("Translation")
    
    if uploaded_file is not None:
        try:
            # Extract content
            content = extract_text_from_docx(uploaded_file)
            paragraphs = [item[1] for item in content if item[0] == "paragraph"]
            
            if not paragraphs:
                st.warning("No text content found in the document.")
                st.stop()
            
            st.info(f"Found {len(paragraphs)} paragraphs to translate")
            
            # Initialize session state
            if 'translating' not in st.session_state:
                st.session_state.translating = False
            if 'translated_paragraphs' not in st.session_state:
                st.session_state.translated_paragraphs = [""] * len(paragraphs)
            if 'current_paragraph' not in st.session_state:
                st.session_state.current_paragraph = 0
            
            # Control buttons
            col_btn1, col_btn2 = st.columns(2)
            
            with col_btn1:
                if st.button("Start Translation", type="primary", disabled=st.session_state.translating):
                    st.session_state.translating = True
                    st.session_state.current_paragraph = 0
                    st.session_state.translated_paragraphs = [""] * len(paragraphs)
                    st.rerun()
            
            with col_btn2:
                if st.button("Stop Translation", disabled=not st.session_state.translating):
                    st.session_state.translating = False
                    st.rerun()
            
            # Live translation process
            if st.session_state.translating and st.session_state.current_paragraph < len(paragraphs):
                current_idx = st.session_state.current_paragraph
                
                # Show progress
                progress = current_idx / len(paragraphs)
                st.progress(progress)
                st.write(f"Translating paragraph {current_idx + 1} of {len(paragraphs)}")
                
                # Translate current paragraph
                with st.spinner(f"Translating paragraph {current_idx + 1}..."):
                    translation = translate_text(
                        paragraphs[current_idx], 
                        INDIC_LANGS[source], 
                        INDIC_LANGS[target],
                        style
                    )
                    st.session_state.translated_paragraphs[current_idx] = translation
                    st.session_state.current_paragraph += 1
                
                # Auto-continue to next paragraph
                if st.session_state.current_paragraph < len(paragraphs):
                    time.sleep(0.5)  # Small delay to show progress
                    st.rerun()
                else:
                    st.session_state.translating = False
                    st.success("Translation completed!")
                    st.rerun()
            
            # Show live editable translations
            if any(t.strip() for t in st.session_state.translated_paragraphs):
                st.subheader("Live Translation - Edit as needed")
                
                edited_translations = []
                for i, (original, translated) in enumerate(zip(paragraphs, st.session_state.translated_paragraphs)):
                    st.write(f"**Paragraph {i+1}:**")
                    
                    # Show original in a disabled text area for reference
                    st.text_area(
                        f"Original {i+1}:",
                        value=original,
                        height=60,
                        disabled=True,
                        key=f"orig_{i}"
                    )
                    
                    # Show editable translation
                    if translated.strip():
                        edited_translation = st.text_area(
                            f"Translation {i+1}:",
                            value=translated,
                            height=100,
                            key=f"trans_{i}",
                            placeholder="Translation will appear here..."
                        )
                    else:
                        edited_translation = st.text_area(
                            f"Translation {i+1}:",
                            value="",
                            height=100,
                            key=f"trans_{i}",
                            placeholder="Waiting for translation..." if st.session_state.translating else "Translation will appear here..."
                        )
                    
                    edited_translations.append(edited_translation)
                    st.markdown("---")
                
                # Download button (only show if translation is complete)
                if not st.session_state.translating and all(t.strip() for t in st.session_state.translated_paragraphs):
                    if st.button("Download Translated Document", type="primary"):
                        # Create new document
                        final_doc = Document()
                        for translation in edited_translations:
                            if translation.strip():
                                final_doc.add_paragraph(translation)
                        
                        # Save to buffer
                        buffer = BytesIO()
                        final_doc.save(buffer)
                        buffer.seek(0)
                        
                        # Download
                        st.download_button(
                            label="Download Word File",
                            data=buffer.getvalue(),
                            file_name=f"translated_{source}_to_{target}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
        
        except Exception as e:
            st.error(f"Error processing document: {str(e)}")
    else:
        st.info("Please upload a Word document to begin translation.")

# Footer
st.markdown("---")
st.markdown(
    """
    **Tips for better translations:**
    - Use 'conversational' or 'casual' style for natural everyday language
    - Review and edit translations for accuracy
    - Consider context when editing technical or domain-specific terms
    """
)
